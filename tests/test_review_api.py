"""Tests for the local, version-aware Design Review API foundation."""

from pathlib import Path
from io import BytesIO
from types import SimpleNamespace

import pytest
from docx import Document as WordDocument
from fastapi.testclient import TestClient
from sqlalchemy import select

from app.domain.evidence import EvidenceChunk
from app.domain.enums import CoverageStatus
from app.domain.models import (
    AnalysisAttempt,
    AnalysisDispatchOutbox,
    AnalysisRun,
    DocumentIndexDispatchOutbox,
    DocumentVersion,
)
from app.api.auth import require_authenticated_user
from app.bootstrap.service_factory import build_design_review_chat_service
from app.repositories import database
from app.services.auth_service import AuthenticatedUser
from app.services.coverage_service import (
    AuditPoint,
    AuditPointJudgment,
    CandidateJudgment,
    CoverageAnalysisService,
)
from app.services.visual_evidence_service import VisualAnalysis, VisualEvidenceService
from app.services.design_review_chat_service import PreparedAnswer
from app.services.analysis_queue import get_analysis_queue
from app.services.analysis_reliability_service import AnalysisReliabilityService
from app.core.config import get_settings


@pytest.fixture
def review_client(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    database.get_engine.cache_clear()
    database.get_session_factory.cache_clear()
    monkeypatch.setenv("DATABASE_URL", f"sqlite:///{tmp_path / 'review.db'}")
    monkeypatch.setenv("ANALYSIS_QUEUE_BACKEND", "memory")
    monkeypatch.setenv("AUTH_SECRET", "test-secret-that-is-long-enough")
    database.initialise_database()
    from app.main import app
    app.dependency_overrides[require_authenticated_user] = lambda: AuthenticatedUser(
        id="engineer-1",
        organization_id="organization-1",
        email="test.engineer@example.com",
        display_name="Test Engineer",
        role="admin",
    )
    with TestClient(app) as client:
        yield client
    app.dependency_overrides.pop(require_authenticated_user, None)
    database.get_engine.cache_clear()
    database.get_session_factory.cache_clear()


def create_design_document(client: TestClient, *, document_type: str = "FS", version: str = "1.0") -> str:
    response = client.post(
        "/documents",
        json={
            "title": "Fleet Manager Functional Specification",
            "document_type": document_type,
            "system": "fleet_manager_wcs",
            "vendor": "Demo Vendor",
            "version": version,
            "status": "draft",
            "file_name": "fleet_manager_fs_v1.0.pdf",
        },
    )
    assert response.status_code == 201
    return response.json()["id"]


def test_document_responses_do_not_expose_storage_paths(review_client: TestClient):
    version_id = create_design_document(review_client)

    documents = review_client.get("/documents")

    assert documents.status_code == 200
    document = next(item for item in documents.json() if item["id"] == version_id)
    assert "storage_path" not in document


def test_evidence_chat_streams_tokens_then_final_citations(review_client: TestClient, monkeypatch: pytest.MonkeyPatch):
    from app.api.routes import chat as chat_routes
    from app.main import app

    evidence = EvidenceChunk(
        chunk_id="chunk-001",
        document_version_id="version-1",
        document_title="Fleet Manager FS",
        document_type="FS",
        version="1.0",
        page=3,
        section="Retention",
        content="Task dispatch records are retained for 90 days.",
        fused_score=0.9,
    )

    class FakeChat:
        def prepare(self, **_):
            return PreparedAnswer("retention period", [evidence], [])

        def stream_answer(self, **_):
            yield "Task dispatch records "
            yield "are retained for 90 days."

    class FakeReviewService:
        def get_review_package(self, _):
            return SimpleNamespace(system="fleet_manager", document_links=[SimpleNamespace(document_version_id="version-1")])

    monkeypatch.setattr(chat_routes, "scoped_review_service", lambda *_: FakeReviewService())
    app.dependency_overrides[build_design_review_chat_service] = lambda: FakeChat()
    try:
        response = review_client.post(
            "/design-review/chat/stream",
            json={"question": "How long are records retained?", "review_package_id": "review-1"},
        )
    finally:
        app.dependency_overrides.pop(build_design_review_chat_service, None)

    assert response.status_code == 200
    assert response.headers["content-type"].startswith("text/event-stream")
    assert 'event: token\ndata: {"text": "Task dispatch records "}' in response.text
    assert 'event: final' in response.text
    assert '"chunk_id": "chunk-001"' in response.text


@pytest.mark.asyncio
async def test_chat_closes_review_session_before_provider_and_stream_body(monkeypatch):
    from app.api.routes import chat as chat_routes
    from app.api.schemas import ReviewChatRequest
    from app.services.design_review_chat_service import GroundedAnswer

    sessions = []

    class FakeSession:
        closed = False

        def close(self):
            self.closed = True

    class FakeReviewService:
        def get_review_package(self, _review_id):
            return SimpleNamespace(
                system="fleet_manager",
                document_links=[SimpleNamespace(document_version_id="document-1")],
            )

    class FakeChat:
        def answer(self, **kwargs):
            assert sessions[-1].closed is True
            assert kwargs["document_version_ids"] == ["document-1"]
            assert kwargs["system"] == "fleet_manager"
            return GroundedAnswer(answer="Ready"), [], "query"

        def prepare(self, **kwargs):
            assert sessions[-1].closed is True
            assert kwargs["document_version_ids"] == ["document-1"]
            assert kwargs["system"] == "fleet_manager"
            return PreparedAnswer(
                "query",
                [],
                [],
                no_evidence_answer=GroundedAnswer(answer="Pending"),
            )

        def stream_answer(self, **_):
            assert sessions[-1].closed is True
            yield "Ready"

    def make_session():
        session = FakeSession()
        sessions.append(session)
        return session

    monkeypatch.setattr(chat_routes, "get_session_factory", lambda: make_session)
    monkeypatch.setattr(
        chat_routes,
        "scoped_review_service",
        lambda _db, _user: FakeReviewService(),
    )
    payload = ReviewChatRequest(question="Ready?", review_package_id="review-1")
    user = AuthenticatedUser(
        id="engineer-1",
        organization_id="organization-1",
        email="engineer@example.com",
        display_name="Engineer",
        role="engineer",
    )

    result = chat_routes.design_review_chat(payload, user, FakeChat())
    stream = chat_routes.stream_design_review_chat(payload, user, FakeChat())

    assert result.answer == "Ready"
    assert len(sessions) == 2
    assert all(session.closed for session in sessions)
    chunks = [chunk async for chunk in stream.body_iterator]
    body = "".join(
        chunk.decode() if isinstance(chunk, bytes) else chunk for chunk in chunks
    )
    assert 'event: final' in body
    assert all(session.closed for session in sessions)


def test_chat_unknown_errors_are_logged_and_hidden_for_http_and_stream(
    review_client: TestClient,
    monkeypatch: pytest.MonkeyPatch,
    caplog: pytest.LogCaptureFixture,
):
    from app.api.routes import chat as chat_routes
    from app.main import app

    class FailingChat:
        def answer(self, **_):
            raise RuntimeError("provider credential leaked by adapter")

        def prepare(self, **_):
            raise RuntimeError("provider credential leaked by stream adapter")

    class FakeReviewService:
        def get_review_package(self, _):
            return SimpleNamespace(
                system="fleet_manager",
                document_links=[SimpleNamespace(document_version_id="version-1")],
            )

    monkeypatch.setattr(chat_routes, "scoped_review_service", lambda *_: FakeReviewService())
    app.dependency_overrides[build_design_review_chat_service] = lambda: FailingChat()
    caplog.set_level("ERROR", logger="industrial_rag.app.api.routes.chat")
    try:
        http_response = review_client.post(
            "/design-review/chat",
            json={"question": "What failed?", "review_package_id": "review-secret-boundary"},
        )
        stream_response = review_client.post(
            "/design-review/chat/stream",
            json={"question": "What failed?", "review_package_id": "review-secret-boundary"},
        )
    finally:
        app.dependency_overrides.pop(build_design_review_chat_service, None)

    assert http_response.status_code == 502
    assert http_response.json()["detail"] == "The evidence answer could not be generated."
    assert "provider credential" not in http_response.text
    assert 'event: error\ndata: {"detail": "The evidence answer could not be generated."}' in stream_response.text
    assert "provider credential" not in stream_response.text
    error_records = [
        record
        for record in caplog.records
        if "review-secret-boundary" in record.getMessage()
    ]
    assert len(error_records) == 2
    assert all(record.exc_info for record in error_records)


def test_document_archive_preserves_auditable_record(review_client: TestClient):
    version_id = create_design_document(review_client)

    response = review_client.post(
        f"/documents/{version_id}/archive",
        json={"reason": "The source file exceeded the permitted upload size."},
    )

    assert response.status_code == 200
    archived = response.json()
    assert archived["status"] == "archived"
    assert archived["archived_reason"] == "The source file exceeded the permitted upload size."
    assert archived["archived_by_user_id"] == "engineer-1"
    assert archived["archived_at"] is not None
    assert any(item["id"] == version_id and item["status"] == "archived" for item in review_client.get("/documents").json())


def test_document_in_frozen_review_package_cannot_be_archived(review_client: TestClient):
    baseline = review_client.post(
        "/requirement-baselines", json={"name": "Archive guard URS", "system": "fleet_manager"}
    ).json()
    review_client.post(
        f"/requirement-baselines/{baseline['id']}/requirements/import",
        files={"file": ("urs.csv", b"requirement_code,requirement_text\nURS-001,Record retention\n", "text/csv")},
    )
    version_id = create_design_document(review_client)
    review_client.post(
        "/review-packages",
        json={
            "name": "Archive guard review",
            "system": "fleet_manager",
            "requirement_baseline_id": baseline["id"],
            "design_document_version_ids": [version_id],
        },
    )

    response = review_client.post(
        f"/documents/{version_id}/archive", json={"reason": "Attempted archive after the scope was frozen."}
    )

    assert response.status_code == 400
    assert "frozen Review Package" in response.json()["detail"]


def test_frozen_document_cannot_be_uploaded_or_reparsed(review_client: TestClient):
    version_id = create_design_document(review_client)
    uploaded = review_client.post(
        f"/documents/{version_id}/upload",
        files={"file": ("source.csv", b"interface,owner\nWCS API,Automation\n", "text/csv")},
    )
    assert uploaded.status_code == 200
    baseline = review_client.post(
        "/requirement-baselines", json={"name": "Ingestion guard URS", "system": "fleet_manager"}
    ).json()
    review_client.post(
        f"/requirement-baselines/{baseline['id']}/requirements/import",
        files={
            "file": (
                "urs.csv",
                b"requirement_code,requirement_text\nURS-001,Record retention\n",
                "text/csv",
            )
        },
    )
    review = review_client.post(
        "/review-packages",
        json={
            "name": "Ingestion guard review",
            "system": "fleet_manager",
            "requirement_baseline_id": baseline["id"],
            "design_document_version_ids": [version_id],
        },
    )
    assert review.status_code == 201

    replacement = review_client.post(
        f"/documents/{version_id}/upload",
        files={"file": ("replacement.csv", b"interface,owner\nMES API,IT\n", "text/csv")},
    )
    reparsed = review_client.post(f"/documents/{version_id}/reparse")

    assert replacement.status_code == 400
    assert reparsed.status_code == 400
    assert "frozen Review Package" in replacement.json()["detail"]
    assert "frozen Review Package" in reparsed.json()["detail"]


@pytest.mark.parametrize("ingestion_status", ["index_queued", "indexing"])
def test_active_index_prevents_upload_and_reparse(
    review_client: TestClient, ingestion_status: str
):
    version_id = create_design_document(review_client)
    uploaded = review_client.post(
        f"/documents/{version_id}/upload",
        files={"file": ("source.csv", b"interface,owner\nWCS API,Automation\n", "text/csv")},
    )
    assert uploaded.status_code == 200
    with database.get_session_factory()() as db:
        version = db.get(DocumentVersion, version_id)
        version.ingestion_status = ingestion_status
        db.commit()

    replacement = review_client.post(
        f"/documents/{version_id}/upload",
        files={"file": ("replacement.csv", b"interface,owner\nMES API,IT\n", "text/csv")},
    )
    reparsed = review_client.post(f"/documents/{version_id}/reparse")

    assert replacement.status_code == 400
    assert reparsed.status_code == 400
    assert "indexing" in replacement.json()["detail"]
    assert "indexing" in reparsed.json()["detail"]


def test_archived_document_cannot_be_added_to_a_new_review_package(review_client: TestClient):
    baseline = review_client.post(
        "/requirement-baselines", json={"name": "Archived source URS", "system": "fleet_manager"}
    ).json()
    review_client.post(
        f"/requirement-baselines/{baseline['id']}/requirements/import",
        files={"file": ("urs.csv", b"requirement_code,requirement_text\nURS-001,Record retention\n", "text/csv")},
    )
    version_id = create_design_document(review_client)
    review_client.post(
        f"/documents/{version_id}/archive", json={"reason": "Superseded before the review package was created."}
    )

    response = review_client.post(
        "/review-packages",
        json={
            "name": "Archived source review",
            "system": "fleet_manager",
            "requirement_baseline_id": baseline["id"],
            "design_document_version_ids": [version_id],
        },
    )

    assert response.status_code == 400
    assert "Archived document versions" in response.json()["detail"]


def test_import_requirements_and_create_review_package(review_client: TestClient):
    baseline = review_client.post(
        "/requirement-baselines",
        json={"name": "Pseudo URS v1.0", "system": "fleet_manager_wcs"},
    )
    assert baseline.status_code == 201
    baseline_id = baseline.json()["id"]

    requirements = review_client.post(
        f"/requirement-baselines/{baseline_id}/requirements/import",
        files={
            "file": (
                "urs.csv",
                b"requirement_code,requirement_text,priority,category\n"
                b"AGV-URS-001,The Fleet Manager shall reserve exclusive zones.,High,Traffic management\n"
                b"AGV-URS-002,The system shall provide an audit history.,High,Auditability\n",
                "text/csv",
            )
        },
    )
    assert requirements.status_code == 200
    assert requirements.json()["imported_count"] == 2

    version_id = create_design_document(review_client)
    review = review_client.post(
        "/review-packages",
        json={
            "name": "DR-001 Fleet Manager Review",
            "system": "fleet_manager_wcs",
            "requirement_baseline_id": baseline_id,
            "design_document_version_ids": [version_id],
        },
    )
    assert review.status_code == 201
    assert review.json()["design_document_version_ids"] == [version_id]
    assert review.json()["requirement_count"] == 2

    analysis = review_client.post(f"/review-packages/{review.json()['id']}/analyses")
    assert analysis.status_code == 202
    assert analysis.json()["status"] == "queued"

    # The HTTP request commits only database work. Redis/RQ publication belongs
    # to analysis-maintenance, so every Outbox row is still pending here.
    with database.get_session_factory()() as db:
        outboxes = list(db.scalars(select(AnalysisDispatchOutbox)))
        assert len(outboxes) == 2
        assert all(item.status == "pending" for item in outboxes)
        assert all(item.job_id is None for item in outboxes)

    resumed_runs = review_client.get(f"/review-packages/{review.json()['id']}/analyses")
    assert resumed_runs.status_code == 200
    assert [run["id"] for run in resumed_runs.json()] == [analysis.json()["id"]]


def test_document_index_submission_returns_accepted_without_running_embeddings(
    review_client: TestClient,
):
    version_id = create_design_document(review_client)
    uploaded = review_client.post(
        f"/documents/{version_id}/upload",
        files={
            "file": (
                "interfaces.csv",
                b"interface,owner\nWCS API,Automation\n",
                "text/csv",
            )
        },
    )
    assert uploaded.status_code == 200
    assert uploaded.json()["ingestion_status"] == "parsed_pending_index"

    queued = review_client.post(f"/documents/{version_id}/index")

    assert queued.status_code == 202
    assert queued.json()["ingestion_status"] == "index_queued"
    assert queued.json()["ingestion_error"] is None
    with database.get_session_factory()() as db:
        index_outbox = db.scalar(
            select(DocumentIndexDispatchOutbox).where(
                DocumentIndexDispatchOutbox.document_version_id == version_id
            )
        )
        assert index_outbox is not None
        assert index_outbox.status == "pending"
        assert index_outbox.job_id is None
    persisted = next(
        item for item in review_client.get("/documents").json() if item["id"] == version_id
    )
    assert persisted["ingestion_status"] == "index_queued"

    duplicate = review_client.post(f"/documents/{version_id}/index")
    assert duplicate.status_code == 409
    assert "already queued or running" in duplicate.json()["detail"]


def test_index_unknown_error_is_logged_with_document_id_and_hidden(
    review_client: TestClient,
    monkeypatch: pytest.MonkeyPatch,
    caplog: pytest.LogCaptureFixture,
):
    from app.services.indexing_service import DocumentIndexSubmissionService

    version_id = create_design_document(review_client)

    def fail_submission(_self, _document_version_id):
        raise RuntimeError("redis://user:secret@internal-host")

    monkeypatch.setattr(
        DocumentIndexSubmissionService,
        "queue_document_version",
        fail_submission,
    )
    caplog.set_level("ERROR", logger="industrial_rag.app.api.routes.documents")

    response = review_client.post(f"/documents/{version_id}/index")

    assert response.status_code == 503
    assert response.json()["detail"] == "Document indexing is temporarily unavailable."
    assert "internal-host" not in response.text
    records = [record for record in caplog.records if version_id in record.getMessage()]
    assert len(records) == 1
    assert records[0].exc_info


def test_visual_analysis_unknown_error_is_logged_with_document_id_and_hidden(
    review_client: TestClient,
    monkeypatch: pytest.MonkeyPatch,
    caplog: pytest.LogCaptureFixture,
):
    from app.api.dependencies import get_visual_evidence_service
    from app.api.routes import documents as document_routes
    from app.bootstrap.service_factory import build_visual_interpreter
    from app.main import app

    class FailingVisualEvidence:
        def analyse_figures(self, _document_version_id, _interpreter):
            raise RuntimeError("vision provider secret")

    app.dependency_overrides[get_visual_evidence_service] = lambda: FailingVisualEvidence()
    app.dependency_overrides[build_visual_interpreter] = lambda: object()
    monkeypatch.setattr(
        document_routes,
        "get_settings",
        lambda: SimpleNamespace(enable_visual_analysis=True),
    )
    caplog.set_level("ERROR", logger="industrial_rag.app.api.routes.documents")
    try:
        response = review_client.post(
            "/documents/document-visual-boundary/figures/analyse"
        )
    finally:
        app.dependency_overrides.pop(get_visual_evidence_service, None)
        app.dependency_overrides.pop(build_visual_interpreter, None)

    assert response.status_code == 502
    assert response.json()["detail"] == "Visual analysis is temporarily unavailable."
    assert "provider secret" not in response.text
    records = [
        record
        for record in caplog.records
        if "document-visual-boundary" in record.getMessage()
    ]
    assert len(records) == 1
    assert records[0].exc_info


def test_progress_reconciles_a_stale_run_after_all_items_settle(review_client: TestClient):
    baseline = review_client.post(
        "/requirement-baselines", json={"name": "Reconciled status URS", "system": "fleet_manager"}
    ).json()
    review_client.post(
        f"/requirement-baselines/{baseline['id']}/requirements/import",
        files={"file": ("urs.csv", b"requirement_code,requirement_text\nURS-001,Record retention\n", "text/csv")},
    )
    review = review_client.post(
        "/review-packages",
        json={
            "name": "Reconciled status review",
            "system": "fleet_manager",
            "requirement_baseline_id": baseline["id"],
            "design_document_version_ids": [create_design_document(review_client)],
        },
    ).json()
    run = review_client.post(f"/review-packages/{review['id']}/analyses").json()

    db = database.get_session_factory()()
    persisted = db.get(AnalysisRun, run["id"])
    assert persisted is not None
    for item in persisted.items:
        item.status = "completed"
    persisted.status = "running"
    persisted.completed_at = None
    db.commit()
    db.close()

    progress = review_client.get(f"/analysis-runs/{run['id']}/progress").json()

    assert progress["status"] == "completed"
    assert progress["completed_items"] == 1


def test_progress_requeues_items_when_their_rq_jobs_are_stale(review_client: TestClient):
    from app.main import app

    class FakeQueue:
        def __init__(self):
            self.enqueued: list[str] = []

        def enqueue_dispatches(self, dispatches):
            values = list(dispatches)
            self.enqueued.extend(dispatch.item_id for dispatch in values)
            return {dispatch.item_id: dispatch.job_id for dispatch in values}

        def stale_item_ids(self, job_ids):
            return {item_id for item_id, job_id in job_ids.items() if job_id.startswith("orphan-")}

    queue = FakeQueue()
    app.dependency_overrides[get_analysis_queue] = lambda: queue
    try:
        baseline = review_client.post(
            "/requirement-baselines", json={"name": "Orphan recovery URS", "system": "fleet_manager"}
        ).json()
        review_client.post(
            f"/requirement-baselines/{baseline['id']}/requirements/import",
            files={"file": ("urs.csv", b"requirement_code,requirement_text\nURS-001,Record retention\n", "text/csv")},
        )
        review = review_client.post(
            "/review-packages",
            json={
                "name": "Orphan recovery review",
                "system": "fleet_manager",
                "requirement_baseline_id": baseline["id"],
                "design_document_version_ids": [create_design_document(review_client)],
            },
        ).json()
        run = review_client.post(f"/review-packages/{review['id']}/analyses").json()

        db = database.get_session_factory()()
        persisted = db.get(AnalysisRun, run["id"])
        assert persisted is not None
        persisted.items[0].job_id = f"orphan-{persisted.items[0].id}"
        db.commit()
        db.close()

        db = database.get_session_factory()()
        result = AnalysisReliabilityService(db, get_settings()).tick(queue)
        db.close()

        progress = review_client.get(f"/analysis-runs/{run['id']}/progress").json()
        assert progress["queued_items"] == 1
        assert result["stale_jobs"] == 1
        assert queue.enqueued.count(progress["items"][0]["id"]) == 2

        db = database.get_session_factory()()
        refreshed = db.get(AnalysisRun, run["id"])
        assert refreshed is not None
        assert refreshed.items[0].dispatch_version == 1
        assert refreshed.items[0].job_id == f"analysis-{refreshed.items[0].id}-1"
        db.close()
    finally:
        app.dependency_overrides.pop(get_analysis_queue, None)


def test_manual_retry_resets_cycle_attempts_and_preserves_attempt_history(
    review_client: TestClient,
):
    baseline = review_client.post(
        "/requirement-baselines",
        json={"name": "Retry cycle URS", "system": "fleet_manager"},
    ).json()
    review_client.post(
        f"/requirement-baselines/{baseline['id']}/requirements/import",
        files={
            "file": (
                "urs.csv",
                b"requirement_code,requirement_text\nURS-001,Record retention\n",
                "text/csv",
            )
        },
    )
    review = review_client.post(
        "/review-packages",
        json={
            "name": "Retry cycle review",
            "system": "fleet_manager",
            "requirement_baseline_id": baseline["id"],
            "design_document_version_ids": [create_design_document(review_client)],
        },
    ).json()
    run = review_client.post(f"/review-packages/{review['id']}/analyses").json()

    with database.get_session_factory()() as db:
        persisted_run = db.get(AnalysisRun, run["id"])
        item = persisted_run.items[0]
        item.status = "failed"
        item.attempt_count = 3
        persisted_run.status = "failed"
        db.add_all(
            [
                AnalysisAttempt(
                    analysis_run_item_id=item.id,
                    attempt_number=attempt_number,
                    worker_id="worker-old-cycle",
                    status="failed",
                )
                for attempt_number in range(1, 4)
            ]
        )
        db.commit()
        item_id = item.id

    retried = review_client.post(f"/analysis-runs/{run['id']}/retry")

    assert retried.status_code == 202
    assert retried.json()["status"] == "queued"
    with database.get_session_factory()() as db:
        item = db.get(AnalysisRun, run["id"]).items[0]
        assert item.attempt_count == 0
        assert item.dispatch_version == 1
        service = CoverageAnalysisService(db, retrieval=None, judge=None)
        claims = [
            service._claim_attempt(
                item_id,
                worker_id="worker-new-cycle",
                max_attempts=3,
                lease_seconds=60,
                expected_dispatch_version=1,
            )
            for _ in range(4)
        ]
        attempts = list(
            db.scalars(
                select(AnalysisAttempt)
                .where(AnalysisAttempt.analysis_run_item_id == item_id)
                .order_by(AnalysisAttempt.attempt_number)
            )
        )

    assert [claim[1:] if claim else None for claim in claims] == [
        (1, 4),
        (2, 5),
        (3, 6),
        None,
    ]
    assert [attempt.attempt_number for attempt in attempts] == [1, 2, 3, 4, 5, 6]


@pytest.mark.asyncio
async def test_sse_uses_short_sessions_in_threadpool_and_reads_poll_interval_once(monkeypatch):
    from app.api.routes import analysis_runs as analysis_routes

    sessions = []
    threadpool_calls = []
    settings_calls = 0

    class FakeSession:
        closed = False

        def close(self):
            self.closed = True

    class FakeService:
        def get_analysis_run(self, _run_id):
            return object()

    class FakeProgress:
        status = "completed"

        def model_dump_json(self):
            return '{"status":"completed"}'

    class FakeRequest:
        async def is_disconnected(self):
            return False

    def make_session():
        session = FakeSession()
        sessions.append(session)
        return session

    async def tracked_threadpool(func, *args):
        threadpool_calls.append(func.__name__)
        return func(*args)

    def fake_settings():
        nonlocal settings_calls
        settings_calls += 1
        return SimpleNamespace(analysis_progress_poll_seconds=99)

    monkeypatch.setattr(analysis_routes, "get_session_factory", lambda: make_session)
    monkeypatch.setattr(analysis_routes, "scoped_review_service", lambda _db, _user: FakeService())
    monkeypatch.setattr(analysis_routes, "analysis_progress_response", lambda _run: FakeProgress())
    monkeypatch.setattr(analysis_routes, "run_in_threadpool", tracked_threadpool)
    monkeypatch.setattr(analysis_routes, "get_settings", fake_settings)
    user = AuthenticatedUser(
        id="engineer-1",
        organization_id="organization-1",
        email="test.engineer@example.com",
        display_name="Test Engineer",
        role="engineer",
    )

    response = await analysis_routes.stream_analysis_run_progress(
        "run-1", FakeRequest(), user
    )

    assert len(sessions) == 1
    assert sessions[0].closed is True
    chunks = [chunk async for chunk in response.body_iterator]
    body = "".join(
        chunk.decode() if isinstance(chunk, bytes) else chunk for chunk in chunks
    )
    assert body == (
        'event: progress\ndata: {"status":"completed"}\n\n'
        'event: complete\ndata: {"status":"completed"}\n\n'
    )
    assert len(sessions) == 2
    assert all(session.closed for session in sessions)
    assert threadpool_calls == ["_load_analysis_progress", "_load_analysis_progress"]
    assert settings_calls == 1


def test_sse_progress_query_closes_its_session_on_error(monkeypatch):
    from app.api.routes import analysis_runs as analysis_routes

    class FakeSession:
        closed = False

        def close(self):
            self.closed = True

    class FailingService:
        def get_analysis_run(self, _run_id):
            raise RuntimeError("database query failed")

    session = FakeSession()
    monkeypatch.setattr(analysis_routes, "get_session_factory", lambda: lambda: session)
    monkeypatch.setattr(
        analysis_routes,
        "scoped_review_service",
        lambda _db, _user: FailingService(),
    )

    with pytest.raises(RuntimeError, match="database query failed"):
        analysis_routes._load_analysis_progress("run-1", object())

    assert session.closed is True


def test_review_packages_are_private_to_the_owner_within_an_organization(review_client: TestClient):
    baseline = review_client.post(
        "/requirement-baselines", json={"name": "Private URS", "system": "fleet_manager_wcs"}
    ).json()
    review_client.post(
        f"/requirement-baselines/{baseline['id']}/requirements/import",
        files={"file": ("urs.csv", b"requirement_code,requirement_text\nURS-001,Private requirement\n", "text/csv")},
    )
    package = review_client.post(
        "/review-packages",
        json={
            "name": "Private review",
            "system": "fleet_manager_wcs",
            "requirement_baseline_id": baseline["id"],
            "design_document_version_ids": [create_design_document(review_client)],
        },
    ).json()

    from app.main import app

    app.dependency_overrides[require_authenticated_user] = lambda: AuthenticatedUser(
        id="engineer-2",
        organization_id="organization-1",
        email="other.engineer@example.com",
        display_name="Other Engineer",
        role="engineer",
    )
    assert review_client.get("/review-packages").json() == []
    assert review_client.get(f"/review-packages/{package['id']}").status_code == 404


def test_import_urs_table_creates_traceable_baseline_without_manual_setup(review_client: TestClient):
    response = review_client.post(
        "/requirement-baselines/import",
        data={"name": "AGV Fleet Manager URS v1.0"},
        files={
            "file": (
                "agv_fleet_manager_urs_v1.0.csv",
                "序号,系统,requirement,reasonal/impact,是否critical\n"
                "1,Fleet Manager,The system shall prevent dispatch to unavailable AGVs.,Patient safety and operational continuity,是\n"
                "2,Fleet Manager,The system shall retain task status history.,Deviation investigation support,否\n".encode("utf-8"),
                "text/csv",
            )
        },
    )

    assert response.status_code == 201
    imported = response.json()
    assert imported["baseline"]["name"] == "AGV Fleet Manager URS v1.0"
    assert imported["baseline"]["system"] == "Fleet Manager"
    assert imported["imported_count"] == 2
    assert imported["requirements"][0] == {
        "id": imported["requirements"][0]["id"],
        "requirement_code": "URS-001",
        "requirement_text": "The system shall prevent dispatch to unavailable AGVs.",
        "source_row": "1",
        "requirement_system": "Fleet Manager",
        "rationale_impact": "Patient safety and operational continuity",
        "is_critical": True,
        "priority": None,
        "category": None,
        "source_section": None,
    }


def test_review_package_rejects_non_design_document(review_client: TestClient):
    baseline = review_client.post(
        "/requirement-baselines",
        json={"name": "Pseudo URS v1.1", "system": "fleet_manager_wcs"},
    ).json()
    review_client.post(
        f"/requirement-baselines/{baseline['id']}/requirements/import",
        files={
            "file": (
                "urs.csv",
                b"requirement_code,requirement_text\n"
                b"AGV-URS-001,The Fleet Manager shall reserve exclusive zones.\n",
                "text/csv",
            )
        },
    )
    version_id = create_design_document(review_client, document_type="TECHNICAL_MANUAL")

    response = review_client.post(
        "/review-packages",
        json={
            "name": "DR-invalid",
            "system": "fleet_manager_wcs",
            "requirement_baseline_id": baseline["id"],
            "design_document_version_ids": [version_id],
        },
    )
    assert response.status_code == 400
    assert "Functional, Software Design, or Hardware Design" in response.json()["detail"]


@pytest.mark.parametrize("document_type", ["FS", "SDS", "HDS"])
def test_review_package_accepts_supported_design_specification_types(
    review_client: TestClient, document_type: str
):
    baseline = review_client.post(
        "/requirement-baselines",
        json={"name": f"{document_type} acceptance URS", "system": "fleet_manager"},
    ).json()
    review_client.post(
        f"/requirement-baselines/{baseline['id']}/requirements/import",
        files={
            "file": (
                "urs.csv",
                b"requirement_code,requirement_text\nURS-001,The system shall retain task status history.\n",
                "text/csv",
            )
        },
    )
    version_id = create_design_document(review_client, document_type=document_type)

    response = review_client.post(
        "/review-packages",
        json={
            "name": f"DR-{document_type}-accepted",
            "system": "fleet_manager_wcs",
            "requirement_baseline_id": baseline["id"],
            "design_document_version_ids": [version_id],
        },
    )

    assert response.status_code == 201


def test_document_replacement_preserves_logical_document(review_client: TestClient):
    v1_id = create_design_document(review_client)
    response = review_client.post(
        "/documents",
        json={
            "title": "Fleet Manager Functional Specification",
            "document_type": "FS",
            "system": "fleet_manager_wcs",
            "vendor": "Demo Vendor",
            "version": "1.1",
            "status": "draft",
            "file_name": "fleet_manager_fs_v1.1.pdf",
            "supersedes_version_id": v1_id,
        },
    )
    assert response.status_code == 201
    v2 = response.json()
    original = review_client.get("/documents").json()
    v1 = next(item for item in original if item["id"] == v1_id)
    assert v2["document_id"] == v1["document_id"]
    assert v1["status"] == "superseded"


def test_pdf_upload_rejects_non_extractable_document(review_client: TestClient, tmp_path: Path):
    import fitz

    version_id = create_design_document(review_client)
    pdf_path = tmp_path / "empty.pdf"
    document = fitz.open()
    document.new_page(width=72, height=72)
    document.save(pdf_path)
    document.close()

    response = review_client.post(
        f"/documents/{version_id}/upload",
        files={"file": ("empty.pdf", pdf_path.read_bytes(), "application/pdf")},
    )
    assert response.status_code == 400
    assert "No extractable text" in response.json()["detail"]


def test_encrypted_pdf_requires_and_accepts_a_one_time_password(review_client: TestClient, tmp_path: Path):
    import fitz

    source = fitz.open()
    page = source.new_page()
    page.insert_text((72, 72), "4.1 Task dispatch\nThe system retains task dispatch records for audit.")
    encrypted_path = tmp_path / "protected.pdf"
    source.save(
        encrypted_path,
        encryption=fitz.PDF_ENCRYPT_AES_256,
        owner_pw="supplier-owner-password",
        user_pw="supplier-password",
    )
    source.close()
    encrypted_pdf = encrypted_path.read_bytes()

    missing_password_version = create_design_document(review_client)
    missing_password = review_client.post(
        f"/documents/{missing_password_version}/upload",
        files={"file": ("protected.pdf", encrypted_pdf, "application/pdf")},
    )
    assert missing_password.status_code == 400
    assert "Enter its password" in missing_password.json()["detail"]

    password_version = create_design_document(review_client)
    accepted = review_client.post(
        f"/documents/{password_version}/upload",
        files={"file": ("protected.pdf", encrypted_pdf, "application/pdf")},
        data={"pdf_password": "supplier-password"},
    )
    assert accepted.status_code == 200
    assert accepted.json()["ingestion_status"] == "parsed_pending_index"


def test_pdf_visual_evidence_is_rendered_before_explicit_analysis(
    review_client: TestClient, monkeypatch: pytest.MonkeyPatch
):
    monkeypatch.setenv("ENABLE_VISUAL_ANALYSIS", "false")
    import fitz

    version_id = create_design_document(review_client)
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "5.1 Interface Diagram\nWCS -> Fleet Manager -> AGV")
    page.draw_rect(fitz.Rect(70, 100, 300, 170))
    source_pdf = document.tobytes()
    document.close()

    upload = review_client.post(
        f"/documents/{version_id}/upload",
        files={"file": ("fleet_manager_interface.pdf", source_pdf, "application/pdf")},
    )
    assert upload.status_code == 200

    figures = review_client.get(f"/documents/{version_id}/figures")
    assert figures.status_code == 200
    assert len(figures.json()) == 1
    figure = figures.json()[0]
    assert figure["page"] == 1
    assert figure["analysis_status"] == "extracted"
    assert figure["candidate_description"] is None

    asset = review_client.get(f"/documents/{version_id}/figures/{figure['id']}/asset")
    assert asset.status_code == 200
    assert asset.headers["content-type"] == "image/png"

    disabled = review_client.post(f"/documents/{version_id}/figures/analyse")
    assert disabled.status_code == 403
    assert "disabled" in disabled.json()["detail"]

    class FakeVisualInterpreter:
        def analyse(self, *, image_path: Path, page: int, section: str | None) -> VisualAnalysis:
            assert image_path.is_file()
            assert page == 1
            return VisualAnalysis(
                diagram_type="interface data flow",
                visible_labels=["WCS", "Fleet Manager", "AGV"],
                candidate_description="The diagram visibly shows an interface sequence between WCS, Fleet Manager, and AGV.",
                candidate_relationships=["WCS → Fleet Manager", "Fleet Manager → AGV"],
                limitations="Arrow semantics require engineering review.",
            )

    db = database.get_session_factory()()
    try:
        analysed = VisualEvidenceService(db, Path("data")).analyse_figures(version_id, FakeVisualInterpreter())
        assert analysed[0].analysis_status == "analysed"
        assert analysed[0].citation_chunk_id is not None
    finally:
        db.close()

    chunks = review_client.get(f"/documents/{version_id}/chunks").json()
    visual_chunk = next(item for item in chunks if item["id"] == analysed[0].citation_chunk_id)
    assert visual_chunk["page"] == 1
    assert "Candidate visual interpretation" in visual_chunk["content"]


def test_optional_visual_candidate_failure_does_not_fail_upload_or_reparse(
    review_client: TestClient, monkeypatch: pytest.MonkeyPatch
):
    import fitz

    def fail_visual_extraction(*args, **kwargs):
        raise RuntimeError("visual renderer unavailable")

    monkeypatch.setattr(
        VisualEvidenceService,
        "extract_pdf_candidates",
        fail_visual_extraction,
    )
    source = fitz.open()
    page = source.new_page()
    page.insert_text((72, 72), "5.1 Interface Diagram\nWCS sends tasks to the Fleet Manager.")
    source_bytes = source.tobytes()
    source.close()
    version_id = create_design_document(review_client)

    uploaded = review_client.post(
        f"/documents/{version_id}/upload",
        files={"file": ("interface.pdf", source_bytes, "application/pdf")},
    )
    reparsed = review_client.post(f"/documents/{version_id}/reparse")

    assert uploaded.status_code == 200
    assert uploaded.json()["ingestion_status"] == "parsed_pending_index"
    assert reparsed.status_code == 200
    assert reparsed.json()["ingestion_status"] == "parsed_pending_index"
    chunks = review_client.get(f"/documents/{version_id}/chunks").json()
    assert any("WCS sends tasks" in chunk["content"] for chunk in chunks)


def test_docx_upload_preserves_heading_and_table_row_sources(review_client: TestClient):
    version_id = create_design_document(review_client)
    document = WordDocument()
    document.add_heading("4.2 Zone reservation", level=1)
    document.add_paragraph("The Fleet Manager shall reserve exclusive zones before vehicle dispatch.")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Event"
    table.rows[0].cells[1].text = "Retention"
    table.rows[1].cells[0].text = "Task status change"
    table.rows[1].cells[1].text = "90 days"
    stream = BytesIO()
    document.save(stream)

    response = review_client.post(
        f"/documents/{version_id}/upload",
        files={
            "file": (
                "fleet_manager_fs.docx",
                stream.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 200
    chunks = review_client.get(f"/documents/{version_id}/chunks").json()
    assert any("4.2 Zone reservation" in (chunk["section"] or "") for chunk in chunks)
    assert any("Retention: 90 days" in chunk["content"] for chunk in chunks)


def test_csv_upload_preserves_rows_as_citable_chunks(review_client: TestClient):
    version_id = create_design_document(review_client, document_type="TECHNICAL_MANUAL")
    response = review_client.post(
        f"/documents/{version_id}/upload",
        files={
            "file": (
                "interface_register.csv",
                b"interface,owner,description\nWCS API,Automation,Dispatches tasks to fleet manager\nMES API,IT,Receives execution status\n",
                "text/csv",
            )
        },
    )
    assert response.status_code == 200
    chunks = review_client.get(f"/documents/{version_id}/chunks").json()
    assert len(chunks) == 2
    assert chunks[0]["section"] == "CSV row 1"
    assert "owner: Automation" in chunks[0]["content"]


def test_upload_parsing_runs_in_threadpool(
    review_client: TestClient, monkeypatch: pytest.MonkeyPatch
):
    from app.api.routes import documents as document_routes

    calls: list[str] = []

    async def tracked_threadpool(func, *args, **kwargs):
        calls.append(func.__name__)
        return func(*args, **kwargs)

    monkeypatch.setattr(document_routes, "run_in_threadpool", tracked_threadpool)
    version_id = create_design_document(review_client)

    response = review_client.post(
        f"/documents/{version_id}/upload",
        files={"file": ("source.csv", b"interface,owner\nWCS API,Automation\n", "text/csv")},
    )

    assert response.status_code == 200
    assert calls == ["upload_and_parse"]


def test_upload_reads_only_one_byte_beyond_configured_limit(
    review_client: TestClient, monkeypatch: pytest.MonkeyPatch
):
    monkeypatch.setenv("MAX_UPLOAD_SIZE_MB", "1")
    version_id = create_design_document(review_client)

    response = review_client.post(
        f"/documents/{version_id}/upload",
        files={"file": ("oversized.csv", b"x" * (1024 * 1024 + 1), "text/csv")},
    )

    assert response.status_code == 413
    assert "File exceeds 1 MB limit" in response.json()["detail"]


def test_both_requirement_csv_imports_reject_files_over_the_shared_limit(
    review_client: TestClient, monkeypatch: pytest.MonkeyPatch
):
    monkeypatch.setenv("MAX_UPLOAD_SIZE_MB", "1")
    baseline = review_client.post(
        "/requirement-baselines",
        json={"name": "Bounded CSV baseline", "system": "fleet_manager"},
    ).json()
    oversized = b"x" * (1024 * 1024 + 1)

    create_response = review_client.post(
        "/requirement-baselines/import",
        files={"file": ("oversized.csv", oversized, "text/csv")},
    )
    existing_response = review_client.post(
        f"/requirement-baselines/{baseline['id']}/requirements/import",
        files={"file": ("oversized.csv", oversized, "text/csv")},
    )

    assert create_response.status_code == 413
    assert existing_response.status_code == 413
    assert create_response.json()["detail"] == "File exceeds 1 MB limit"
    assert existing_response.json()["detail"] == "File exceeds 1 MB limit"


def test_requirement_csv_import_routes_are_sync_and_use_bounded_file_reads(monkeypatch):
    import inspect
    from app.api.routes import requirements as requirement_routes

    read_sizes = []

    class FakeFile:
        def read(self, size):
            read_sizes.append(size)
            return b"requirement_code,requirement_text\nURS-001,Ready\n"

    monkeypatch.setattr(
        requirement_routes,
        "get_settings",
        lambda: SimpleNamespace(max_upload_size_mb=1),
    )

    content = requirement_routes.read_csv_upload(
        SimpleNamespace(file=FakeFile())
    )

    assert content.endswith(b"URS-001,Ready\n")
    assert read_sizes == [1024 * 1024 + 1]
    assert not inspect.iscoroutinefunction(
        requirement_routes.import_requirement_baseline
    )
    assert not inspect.iscoroutinefunction(requirement_routes.import_requirements)


@pytest.mark.asyncio
async def test_upload_file_read_is_bounded_to_limit_plus_one(monkeypatch: pytest.MonkeyPatch):
    from fastapi import HTTPException
    from app.api.routes import documents as document_routes

    read_sizes: list[int] = []

    class FakeUpload:
        filename = "oversized.csv"

        async def read(self, size: int):
            read_sizes.append(size)
            return b"x" * size

    monkeypatch.setattr(
        document_routes,
        "get_settings",
        lambda: SimpleNamespace(max_upload_size_mb=1),
    )

    with pytest.raises(HTTPException) as error:
        await document_routes.upload_document(
            "document-version-1",
            ingestion=object(),  # type: ignore[arg-type]
            file=FakeUpload(),  # type: ignore[arg-type]
            pdf_password=None,
        )

    assert error.value.status_code == 413
    assert read_sizes == [1024 * 1024 + 1]


def test_citation_reader_returns_the_requested_passage_in_source_order(review_client: TestClient):
    version_id = create_design_document(review_client, document_type="TECHNICAL_MANUAL")
    response = review_client.post(
        f"/documents/{version_id}/upload",
        files={
            "file": (
                "interface_register.csv",
                b"interface,owner\nWCS API,Automation\nMES API,IT\n",
                "text/csv",
            )
        },
    )
    assert response.status_code == 200
    chunks = review_client.get(f"/documents/{version_id}/chunks").json()

    context = review_client.get(f"/documents/{version_id}/chunks/{chunks[1]['id']}/context")

    assert context.status_code == 200
    payload = context.json()
    assert payload["requested_chunk_id"] == chunks[1]["id"]
    assert [item["id"] for item in payload["chunks"]] == [item["id"] for item in chunks]
    assert "MES API" in payload["chunks"][-1]["content"]


def test_original_pdf_source_is_available_for_the_in_app_page_viewer(review_client: TestClient):
    import fitz

    source = fitz.open()
    page = source.new_page()
    page.insert_text((72, 72), "Fleet Manager source page")
    source_bytes = source.tobytes()
    source.close()
    version_id = create_design_document(review_client)

    uploaded = review_client.post(
        f"/documents/{version_id}/upload",
        files={"file": ("fleet_manager.pdf", source_bytes, "application/pdf")},
    )
    assert uploaded.status_code == 200

    response = review_client.get(f"/documents/{version_id}/source")

    assert response.status_code == 200
    assert response.headers["content-type"].startswith("application/pdf")
    assert response.headers["content-disposition"] == "inline"
    assert response.content == source_bytes


def test_stored_document_can_be_reparsed_with_the_current_chunk_configuration(review_client: TestClient):
    import fitz

    source = fitz.open()
    page = source.new_page()
    source_text = "4.1 Task dispatch\n" + ("The system shall retain task dispatch records for audit. " * 12)
    page.insert_textbox((72, 72, 500, 700), source_text)
    content = source.tobytes()
    source.close()

    version_id = create_design_document(review_client)
    uploaded = review_client.post(
        f"/documents/{version_id}/upload",
        files={"file": ("fleet_manager.pdf", content, "application/pdf")},
    )
    assert uploaded.status_code == 200
    assert uploaded.json()["chunk_count"] == 1

    reparsed = review_client.post(f"/documents/{version_id}/reparse")

    assert reparsed.status_code == 200
    assert reparsed.json()["ingestion_status"] == "parsed_pending_index"
    assert reparsed.json()["chunk_count"] == 1
    chunks = review_client.get(f"/documents/{version_id}/chunks").json()
    assert all(chunk["page"] == 1 for chunk in chunks)


def test_document_upload_rejects_unsupported_source(review_client: TestClient):
    version_id = create_design_document(review_client)
    response = review_client.post(
        f"/documents/{version_id}/upload",
        files={"file": ("supplier_note.txt", b"not a supported file", "text/plain")},
    )
    assert response.status_code == 400
    assert "PDF, DOCX, and CSV" in response.json()["detail"]


def test_coverage_analysis_persists_only_retrieved_evidence(review_client: TestClient):
    class FakeRetrieval:
        def retrieve(self, query, filters, limit=8):
            return [
                EvidenceChunk(
                    chunk_id="chunk-001",
                    document_version_id=filters.document_version_ids[0],
                    document_title="Fleet Manager Functional Specification",
                    document_type="FS",
                    version="1.0",
                    page=7,
                    section="4.1 Zone reservation",
                    content="The Fleet Manager reserves exclusive zones before dispatching a vehicle.",
                )
            ]

    class FakeJudge:
        def judge(self, *, requirement_code, requirement_text, evidence):
            return CandidateJudgment(
                design_status=CoverageStatus.COVERED,
                evidence_chunk_ids=["chunk-001", "fabricated-id"],
                rationale="The selected evidence explicitly describes exclusive-zone reservation.",
            )

    baseline = review_client.post(
        "/requirement-baselines",
        json={"name": "Coverage URS v1.0", "system": "fleet_manager_wcs"},
    ).json()
    review_client.post(
        f"/requirement-baselines/{baseline['id']}/requirements/import",
        files={
            "file": (
                "urs.csv",
                b"requirement_code,requirement_text\nAGV-URS-001,The Fleet Manager shall reserve exclusive zones.\n",
                "text/csv",
            )
        },
    )
    version_id = create_design_document(review_client)
    review = review_client.post(
        "/review-packages",
        json={
            "name": "DR-coverage",
            "system": "fleet_manager_wcs",
            "requirement_baseline_id": baseline["id"],
            "design_document_version_ids": [version_id],
        },
    ).json()
    run = review_client.post(f"/review-packages/{review['id']}/analyses").json()

    db = database.get_session_factory()()
    try:
        CoverageAnalysisService(db, FakeRetrieval(), FakeJudge()).execute(run["id"])
    finally:
        db.close()

    findings = review_client.get(f"/analysis-runs/{run['id']}/findings")
    assert findings.status_code == 200
    result = findings.json()[0]
    assert result["design_status"] == "covered"
    assert [evidence["chunk_id"] for evidence in result["evidence"]] == ["chunk-001"]
    progress = review_client.get(f"/analysis-runs/{run['id']}/progress")
    assert progress.status_code == 200
    assert progress.json()["completed_items"] == 1
    assert progress.json()["items"][0]["status"] == "completed"


def test_matrix_keeps_one_row_for_every_requirement_including_unfinished_items(review_client: TestClient):
    baseline = review_client.post(
        "/requirement-baselines",
        json={"name": "Matrix completeness URS", "system": "fleet_manager_wcs"},
    ).json()
    review_client.post(
        f"/requirement-baselines/{baseline['id']}/requirements/import",
        files={
            "file": (
                "urs.csv",
                b"requirement_code,requirement_text,priority\n"
                b"AGV-URS-001,The Fleet Manager shall reserve exclusive zones.,High\n"
                b"AGV-URS-002,The Fleet Manager shall retain audit history.,Medium\n",
                "text/csv",
            )
        },
    )
    version_id = create_design_document(review_client)
    review = review_client.post(
        "/review-packages",
        json={
            "name": "DR-matrix-completeness",
            "system": "fleet_manager_wcs",
            "requirement_baseline_id": baseline["id"],
            "design_document_version_ids": [version_id],
        },
    ).json()
    run = review_client.post(f"/review-packages/{review['id']}/analyses").json()

    matrix = review_client.get(f"/analysis-runs/{run['id']}/matrix")
    assert matrix.status_code == 200
    rows = matrix.json()
    assert [row["requirement_code"] for row in rows] == ["AGV-URS-001", "AGV-URS-002"]
    assert all(row["analysis_status"] == "queued" for row in rows)
    assert all(row["design_status"] is None for row in rows)


def test_composite_requirement_cannot_be_marked_covered_when_one_audit_point_is_not_covered(
    review_client: TestClient,
):
    class FakeRetrieval:
        def retrieve(self, query, filters, limit=8):
            chunk_id = "chunk-interface" if "interface" in query.lower() else "chunk-audit"
            return [
                EvidenceChunk(
                    chunk_id=chunk_id,
                    document_version_id=filters.document_version_ids[0],
                    document_title="Fleet Manager Functional Specification",
                    document_type="FS",
                    version="1.0",
                    page=7,
                    section="4.1 Integration",
                    content="The FS describes the requested capability.",
                )
            ]

    class FakeJudge:
        def decompose(self, *, requirement_code, requirement_text):
            return [
                AuditPoint(point_id="p1", source_excerpt="provide an interface", review_point="provide an interface"),
                AuditPoint(point_id="p2", source_excerpt="retain audit history", review_point="retain audit history"),
            ]

        def judge(self, *, requirement_code, requirement_text, evidence, audit_points):
            return CandidateJudgment(
                design_status=CoverageStatus.COVERED,
                evidence_chunk_ids=["chunk-interface", "chunk-audit"],
                rationale="The requirement appears covered.",
                audit_points=[
                    AuditPointJudgment(
                        **audit_points[0].model_dump(),
                        design_status=CoverageStatus.COVERED,
                        evidence_chunk_ids=["chunk-interface"],
                        rationale="Interface evidence is explicit.",
                    ),
                    AuditPointJudgment(
                        **audit_points[1].model_dump(),
                        design_status=CoverageStatus.PARTIALLY_COVERED,
                        evidence_chunk_ids=["chunk-audit"],
                        rationale="Retention duration is not stated.",
                    ),
                ],
            )

    baseline = review_client.post(
        "/requirement-baselines",
        json={"name": "Composite URS", "system": "fleet_manager_wcs"},
    ).json()
    review_client.post(
        f"/requirement-baselines/{baseline['id']}/requirements/import",
        files={
            "file": (
                "urs.csv",
                b"requirement_code,requirement_text\n"
                b"AGV-URS-003,The system shall provide an interface and retain audit history.\n",
                "text/csv",
            )
        },
    )
    version_id = create_design_document(review_client)
    review = review_client.post(
        "/review-packages",
        json={
            "name": "DR-composite",
            "system": "fleet_manager_wcs",
            "requirement_baseline_id": baseline["id"],
            "design_document_version_ids": [version_id],
        },
    ).json()
    run = review_client.post(f"/review-packages/{review['id']}/analyses").json()
    db = database.get_session_factory()()
    try:
        CoverageAnalysisService(db, FakeRetrieval(), FakeJudge()).execute(run["id"])
    finally:
        db.close()

    matrix = review_client.get(f"/analysis-runs/{run['id']}/matrix").json()
    assert matrix[0]["design_status"] == "review_required"
    assert len(matrix[0]["audit_points"]) == 2
    assert matrix[0]["audit_points"][1]["design_status"] == "partially_covered"


def test_matrix_excel_export_contains_all_rows_and_report_metadata(review_client: TestClient):
    from openpyxl import load_workbook

    baseline = review_client.post(
        "/requirement-baselines",
        json={"name": "Export URS", "system": "fleet_manager_wcs"},
    ).json()
    review_client.post(
        f"/requirement-baselines/{baseline['id']}/requirements/import",
        files={
            "file": (
                "urs.csv",
                b"requirement_code,requirement_text\nAGV-URS-010,The system shall retain audit history.\n",
                "text/csv",
            )
        },
    )
    version_id = create_design_document(review_client)
    review = review_client.post(
        "/review-packages",
        json={
            "name": "DR-export",
            "system": "fleet_manager_wcs",
            "requirement_baseline_id": baseline["id"],
            "design_document_version_ids": [version_id],
        },
    ).json()
    run = review_client.post(f"/review-packages/{review['id']}/analyses").json()

    response = review_client.get(f"/analysis-runs/{run['id']}/export.xlsx")
    assert response.status_code == 200
    workbook = load_workbook(BytesIO(response.content))
    assert workbook.sheetnames == ["URS Traceability Matrix", "Report Metadata"]
    matrix = workbook["URS Traceability Matrix"]
    assert matrix["A2"].value == "AGV-URS-010"
    assert matrix["F2"].value == "Assessment incomplete / technical exception"


def test_original_strategy_is_persisted_and_retrieves_the_unchanged_urs_once(
    review_client: TestClient,
):
    requirement_text = "The Fleet Manager shall retain the original dispatch record exactly."
    queries: list[tuple[str, int]] = []

    class FakeRetrieval:
        def retrieve(self, query, filters, limit=8):
            queries.append((query, limit))
            return [
                EvidenceChunk(
                    chunk_id="chunk-original",
                    document_version_id=filters.document_version_ids[0],
                    document_title="Fleet Manager Functional Specification",
                    document_type="FS",
                    version="1.0",
                    page=9,
                    section="5.2 Records",
                    content="The Fleet Manager retains the original dispatch record.",
                )
            ]

    class FakeJudge:
        def judge(self, *, requirement_code, requirement_text, evidence, audit_points):
            assert len(audit_points) == 1
            assert audit_points[0].review_point == requirement_text
            return CandidateJudgment(
                design_status=CoverageStatus.COVERED,
                evidence_chunk_ids=["chunk-original"],
                rationale="The record retention is explicit.",
                audit_points=[
                    AuditPointJudgment(
                        **audit_points[0].model_dump(),
                        design_status=CoverageStatus.COVERED,
                        evidence_chunk_ids=["chunk-original"],
                        rationale="The record retention is explicit.",
                    )
                ],
            )

    baseline = review_client.post(
        "/requirement-baselines", json={"name": "Original baseline", "system": "fleet_manager_wcs"}
    ).json()
    review_client.post(
        f"/requirement-baselines/{baseline['id']}/requirements/import",
        files={
            "file": (
                "urs.csv",
                f"requirement_code,requirement_text\nURS-001,{requirement_text}\n".encode(),
                "text/csv",
            )
        },
    )
    version_id = create_design_document(review_client)
    review = review_client.post(
        "/review-packages",
        json={
            "name": "Original workflow review",
            "system": "fleet_manager_wcs",
            "requirement_baseline_id": baseline["id"],
            "design_document_version_ids": [version_id],
        },
    ).json()

    created = review_client.post(
        f"/review-packages/{review['id']}/analyses", json={"strategy": "original"}
    )
    assert created.status_code == 202
    run = created.json()
    assert run["strategy"] == "original"
    assert run["strategy_version"] == "original-v1"
    assert review_client.get(f"/analysis-runs/{run['id']}").json()["strategy"] == "original"

    db = database.get_session_factory()()
    try:
        CoverageAnalysisService(db, FakeRetrieval(), FakeJudge()).execute(run["id"])
        persisted = db.get(AnalysisRun, run["id"])
        trace = persisted.items[0].analysis_trace
    finally:
        db.close()

    assert queries == [(requirement_text, 6)]
    assert trace["strategy"] == "original"
    assert trace["retrieval"]["query_policy"] == "original_urs_byte_for_byte"
    assert trace["retrieval"]["queries"][0]["query"] == requirement_text
    assert trace["retrieval"]["queries"][0]["ranked_chunk_ids"] == ["chunk-original"]


def test_decomposed_strategy_keeps_at_most_three_unique_source_grounded_points():
    requirement_text = (
        "The AMR shall enter a safe state, report the fault status, and retain the event record."
    )

    class Planner:
        def decompose(self, *, requirement_code, requirement_text):
            return [
                AuditPoint(point_id="p1", source_excerpt="enter a safe state", review_point="The AMR shall enter a safe state"),
                AuditPoint(point_id="p2", source_excerpt="report the fault status", review_point="The AMR shall report the fault status"),
                AuditPoint(point_id="p3", source_excerpt="retain the event record", review_point="The AMR shall retain the event record"),
                AuditPoint(point_id="p4", source_excerpt="invented standard", review_point="Meet an invented standard"),
            ]

        def judge(self, **kwargs):
            raise AssertionError("judge is not used in this unit test")

    service = CoverageAnalysisService(None, None, Planner())  # type: ignore[arg-type]
    points = service._audit_points("URS-001", requirement_text)

    assert len(points) == 3
    assert [point.point_id for point in points] == ["p1", "p2", "p3"]
    assert all(point.source_excerpt in requirement_text for point in points)
